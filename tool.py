import os
import requests
from datetime import datetime
from pydantic import BaseModel, Field
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re


class Tools:
    def __init__(self):
        pass

    # Add your custom tools using pure Python code here, make sure to add type hints and descriptions

    def get_user_name_and_email_and_id(self, __user__: dict = {}) -> str:
        """
        Get the user name, Email and ID from the user object.
        """

        # Do not include a descrption for __user__ as it should not be shown in the tool's specification
        # The session user object will be passed as a parameter when the function is called

        print(__user__)
        result = ""

        if "name" in __user__:
            result += f"User: {__user__['name']}"
        if "id" in __user__:
            result += f" (ID: {__user__['id']})"
        if "email" in __user__:
            result += f" (Email: {__user__['email']})"

        if result == "":
            result = "User: Unknown"

        return result

    def get_current_time(self) -> str:
        """
        Get the current time in a more human-readable format.
        """

        now = datetime.now()
        current_time = now.strftime("%I:%M:%S %p")  # Using 12-hour format with AM/PM
        current_date = now.strftime(
            "%A, %B %d, %Y"
        )  # Full weekday, month name, day, and year

        return f"Current Date and Time = {current_date}, {current_time}"

    def calculator(
        self,
        equation: str = Field(
            ..., description="The mathematical equation to calculate."
        ),
    ) -> str:
        """
        Calculate the result of an equation.
        """

        # Avoid using eval in production code
        # https://nedbatchelder.com/blog/201206/eval_really_is_dangerous.html
        try:
            result = eval(equation)
            return f"{equation} = {result}"
        except Exception as e:
            print(e)
            return "Invalid equation"

    def get_current_weather(
        self,
        city: str = Field(
            "New York, NY", description="Get the current weather for a given city."
        ),
    ) -> str:
        """
        Get the current weather for a given city.
        """

        api_key = os.getenv("OPENWEATHER_API_KEY")
        if not api_key:
            return (
                "API key is not set in the environment variable 'OPENWEATHER_API_KEY'."
            )

        base_url = "http://api.openweathermap.org/data/2.5/weather"
        params = {
            "q": city,
            "appid": api_key,
            "units": "metric",  # Optional: Use 'imperial' for Fahrenheit
        }

        try:
            response = requests.get(base_url, params=params)
            response.raise_for_status()  # Raise HTTPError for bad responses (4xx and 5xx)
            data = response.json()

            if data.get("cod") != 200:
                return f"Error fetching weather data: {data.get('message')}"

            weather_description = data["weather"][0]["description"]
            temperature = data["main"]["temp"]
            humidity = data["main"]["humidity"]
            wind_speed = data["wind"]["speed"]

            return f"Weather in {city}: {temperature}°C"
        except requests.RequestException as e:
            return f"Error fetching weather data: {str(e)}"

    def generate_word_document(
        self,
        document_text: str = Field(
            ..., description="The text content to save in the Word document."
        ),
        document_title: str = Field(
            "Untitled Document", description="The title of the document."
        ),
        file_name: str = Field(
            "",
            description="Optional custom filename (without extension). If not provided, auto-generated based on title and timestamp.",
        ),
        __user__: dict = {},
    ) -> str:
        """
        Generate and save a Word document with the provided text content.
        The document will be saved on the server and a download link will be provided.
        """

        try:
            # Save to static directory for web access
            # Try different possible paths
            static_paths = [
                "static/generated_docs",
                "backend/open_webui/static/generated_docs",
                "open_webui/static/generated_docs",
                "/app/backend/open_webui/static/generated_docs"
            ]
            
            upload_dir = None
            for path in static_paths:
                try:
                    os.makedirs(path, exist_ok=True)
                    # Test if we can write to this directory
                    test_file = os.path.join(path, ".test")
                    with open(test_file, 'w') as f:
                        f.write("test")
                    os.remove(test_file)
                    upload_dir = path
                    break
                except:
                    continue
            
            if not upload_dir:
                # Fallback to data/uploads
                upload_dir = "data/uploads/lna"
                os.makedirs(upload_dir, exist_ok=True)

            # Generate filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            if file_name:
                # Sanitize filename
                safe_filename = re.sub(r"[^\w\s-]", "", file_name).strip()
                safe_filename = re.sub(r"[-\s]+", "_", safe_filename)
            else:
                # Generate from title
                safe_title = re.sub(r"[^\w\s-]", "", document_title).strip()
                safe_title = re.sub(r"[-\s]+", "_", safe_title)
                safe_filename = f"{safe_title}_{timestamp}"

            full_filename = f"{safe_filename}.docx"
            file_path = os.path.join(upload_dir, full_filename)

            # Create Word document
            doc = Document()

            # Add title
            title_paragraph = doc.add_paragraph()
            title_run = title_paragraph.add_run(document_title)
            title_run.bold = True
            title_run.font.size = Pt(16)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add spacing after title
            doc.add_paragraph()

            # Process and add content
            # Split by paragraphs and handle markdown-like formatting
            paragraphs = document_text.split("\n\n")

            for para_text in paragraphs:
                if not para_text.strip():
                    continue

                # Check if it's a heading (starts with #)
                if para_text.strip().startswith("#"):
                    heading_text = para_text.strip().lstrip("#").strip()
                    heading_level = len(para_text.strip()) - len(
                        para_text.strip().lstrip("#")
                    )
                    heading_level = min(heading_level, 3)  # Max heading level 3
                    doc.add_heading(heading_text, level=heading_level)
                else:
                    # Regular paragraph
                    # Handle line breaks within paragraph
                    lines = para_text.split("\n")
                    p = doc.add_paragraph()

                    for i, line in enumerate(lines):
                        if i > 0:
                            p.add_run("\n")

                        # Simple formatting: **bold**, *italic*
                        # This is a simplified approach
                        line = line.strip()
                        if line:
                            # Handle bold (**text**)
                            parts = re.split(r"(\*\*.*?\*\*)", line)
                            for part in parts:
                                if part.startswith("**") and part.endswith("**"):
                                    run = p.add_run(part[2:-2])
                                    run.bold = True
                                else:
                                    # Handle italic (*text*)
                                    italic_parts = re.split(r"(\*.*?\*)", part)
                                    for italic_part in italic_parts:
                                        if (
                                            italic_part.startswith("*")
                                            and italic_part.endswith("*")
                                            and not italic_part.startswith("**")
                                        ):
                                            run = p.add_run(italic_part[1:-1])
                                            run.italic = True
                                        else:
                                            p.add_run(italic_part)

            # Add metadata as footer (optional)
            section = doc.sections[0]
            footer = section.footer
            footer_para = (
                footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            )
            footer_para.text = (
                f"Generated on {datetime.now().strftime('%d.%m.%Y %H:%M')}"
            )
            footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            if __user__ and "name" in __user__:
                footer_para.text += f" | User: {__user__['name']}"

            # Save document
            doc.save(file_path)

            # Generate download URL
            if "static" in upload_dir:
                # File is in static directory, can be accessed via web
                download_url = f"/static/generated_docs/{full_filename}"
                return f"Документ успешно создан и сохранен!\n\n📄 Имя файла: {full_filename}\n\n🔗 Ссылка для скачивания:\n{download_url}\n\nПопробуйте открыть эту ссылку в браузере. Если не работает, файл находится на сервере по пути: {file_path}"
            else:
                # File is in data directory
                return f"Документ успешно создан и сохранен!\n\n📄 Имя файла: {full_filename}\n📁 Путь на сервере: {file_path}\n\nДокумент сохранен на сервере. Для настройки скачивания через веб-интерфейс обратитесь к администратору для настройки nginx."

        except Exception as e:
            return f"Error generating Word document: {str(e)}\n\nPlease ensure the python-docx library is installed on the server."
